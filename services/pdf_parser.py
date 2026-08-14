import re
import io
from pypdf import PdfReader
import docx

def extract_text_from_file(file_stream, filename="resume.pdf"):
    """
    Unified text extractor for both PDF and DOCX resume files.
    """
    ext = filename.lower().split('.')[-1]
    
    if ext == 'docx':
        return extract_text_from_docx(file_stream)
    else:
        return extract_text_from_pdf(file_stream)

def extract_text_from_pdf(pdf_file_stream):
    """
    Extracts raw text from a PDF file stream.
    """
    try:
        reader = PdfReader(pdf_file_stream)
        text_content = []
        page_count = len(reader.pages)
        
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
                
        full_text = "\n".join(text_content)
        words = re.findall(r'\b\w+\b', full_text)
        word_count = len(words)
        
        sections = parse_resume_sections(full_text)
        
        return {
            'success': True,
            'raw_text': full_text,
            'page_count': page_count,
            'word_count': word_count,
            'sections': sections
        }
    except Exception as e:
        return {
            'success': False,
            'error': f"PDF Extraction Error: {str(e)}",
            'raw_text': '',
            'page_count': 0,
            'word_count': 0,
            'sections': {}
        }

def extract_text_from_docx(docx_file_stream):
    """
    Extracts raw text from a Microsoft Word (.docx) file stream.
    """
    try:
        doc = docx.Document(docx_file_stream)
        full_text_list = []
        
        # Read paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                full_text_list.append(p.text.strip())
                
        # Read table cell content if present
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    full_text_list.append(row_text)
                    
        full_text = "\n".join(full_text_list)
        words = re.findall(r'\b\w+\b', full_text)
        word_count = len(words)
        
        sections = parse_resume_sections(full_text)
        
        return {
            'success': True,
            'raw_text': full_text,
            'page_count': 1,
            'word_count': word_count,
            'sections': sections
        }
    except Exception as e:
        return {
            'success': False,
            'error': f"DOCX Extraction Error: {str(e)}",
            'raw_text': '',
            'page_count': 0,
            'word_count': 0,
            'sections': {}
        }

def parse_resume_sections(text):
    """
    Identifies common resume section headers and extracts text under each section.
    """
    section_headers = {
        'summary': r'(?:summary|objective|profile|about me)',
        'experience': r'(?:work experience|professional experience|employment history|experience)',
        'education': r'(?:education|academic background|qualifications)',
        'skills': r'(?:skills|technical skills|technologies|core competencies)',
        'projects': r'(?:projects|personal projects|key projects)',
        'certifications': r'(?:certifications|licenses|courses)'
    }
    
    sections = {}
    lines = text.split('\n')
    current_section = 'general'
    sections[current_section] = []
    
    for line in lines:
        clean_line = line.strip().lower()
        matched_header = False
        for sec_name, pattern in section_headers.items():
            if re.match(r'^' + pattern + r'[:\s]*$', clean_line, re.IGNORECASE):
                current_section = sec_name
                if current_section not in sections:
                    sections[current_section] = []
                matched_header = True
                break
        if not matched_header:
            sections[current_section].append(line)
            
    return {k: "\n".join(v).strip() for k, v in sections.items() if v}
