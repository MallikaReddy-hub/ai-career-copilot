import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import io
import re

def create_resume_docx(resume_text, target_title="Software Engineer"):
    """
    Creates a clean, executive ATS-compliant Microsoft Word (.docx) document.
    """
    doc = docx.Document()

    # Set standard 0.65-inch margins for executive ATS layout
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    # Style definitions
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(30, 41, 59)

    known_headers = [
        'PROFILE', 'SUMMARY', 'PROFESSIONAL SUMMARY', 'OBJECTIVE',
        'TECHNICAL SKILLS', 'SKILLS', 'CORE COMPETENCIES', 'TECHNOLOGIES',
        'EXPERIENCE', 'PROFESSIONAL EXPERIENCE', 'WORK EXPERIENCE', 'EMPLOYMENT HISTORY',
        'PROJECTS', 'KEY PROJECTS', 'ACADEMIC PROJECTS',
        'EDUCATION', 'ACADEMIC BACKGROUND', 'CERTIFICATIONS', 'ACHIEVEMENTS'
    ]

    lines = resume_text.strip().split('\n')
    is_header = True

    for i, line in enumerate(lines):
        trimmed = line.strip()
        if not trimmed:
            continue

        clean_upper = trimmed.upper().rstrip(':')

        # Header candidate name (first line)
        if is_header and (i == 0 or (len(trimmed.split()) <= 5 and '@' not in trimmed and 'http' not in trimmed)):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(trimmed)
            run.font.size = Pt(18)
            run.bold = True
            run.font.color.rgb = RGBColor(15, 23, 42)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            continue
        elif is_header and ('@' in trimmed or 'linkedin' in trimmed.lower() or 'github' in trimmed.lower() or '|' in trimmed or '+' in trimmed):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(trimmed)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 116, 139)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(10)
            is_header = False
            continue

        is_header = False

        # Section Header Detection
        if clean_upper in known_headers:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(clean_upper)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(15, 23, 42)
            
            # Add subtle horizontal line below heading
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                             r'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="CBD5E1"/>'
                             r'</w:pBdr>')
            pPr.append(pBdr)
            continue

        # Bullet point detection
        if trimmed.startswith('•') or trimmed.startswith('-') or trimmed.startswith('*'):
            bullet_content = re.sub(r'^[•\-\*\d\.]+\s*', '', trimmed).strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(bullet_content)
            continue

        # Field bolding (e.g. "Languages: Python, Java...")
        if ':' in trimmed and len(trimmed.split(':')[0].split()) <= 4:
            parts = trimmed.split(':', 1)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            r1 = p.add_run(parts[0].strip() + ": ")
            r1.bold = True
            r2 = p.add_run(parts[1].strip())
            continue

        # Standard paragraph text
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(trimmed)

    # Save to in-memory bytes buffer
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
